"""
文档解析模块 - 保留完整结构
支持：PDF, EPUB, FB2, TXT, DOCX, HTML
"""

import fitz  # PyMuPDF
from ebooklib import epub
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field
import re
import base64
import json


@dataclass
class Chapter:
    """章节数据结构"""
    id: int
    title: str
    content: str  # 原始 XHTML/HTML 内容
    text: str    # 纯文本（用于翻译）
    images: List[Dict] = field(default_factory=list)  # 图片信息
    start_pos: int = 0
    end_pos: int = 0


@dataclass
class Document:
    """文档数据结构"""
    title: str
    source_type: str  # pdf/epub/fb2/txt/docx/html
    chapters: List[Chapter] = field(default_factory=list)
    images: List[Dict] = field(default_factory=list)  # 所有图片
    metadata: Dict = field(default_factory=dict)
    css: str = ""  # 样式表
    original_path: str = ""


async def parse_document(file_path: str, output_dir: str) -> Document:
    """解析文档入口"""
    file_path = Path(file_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    suffix = file_path.suffix.lower()
    
    # 格式映射
    format_map = {
        '.pdf': 'pdf',
        '.epub': 'epub',
        '.fb2': 'fb2',
        '.txt': 'txt',
        '.docx': 'docx',
        '.html': 'html',
        '.htm': 'html',
    }
    
    if suffix not in format_map:
        raise ValueError(f"不支持的格式：{suffix}")
    
    source_type = format_map[suffix]
    
    # 调用对应解析器
    parser_func = PARSER_MAP.get(source_type)
    if not parser_func:
        raise ValueError(f"未实现的解析器：{source_type}")
    
    doc = await parser_func(file_path, output_dir)
    doc.original_path = str(file_path)
    
    return doc


async def _parse_epub(file_path: Path, output_dir: Path) -> Document:
    """解析 EPUB - 保留完整结构"""
    book = epub.read_epub(file_path)
    document = Document(title=file_path.stem, source_type='epub')
    
    # 提取元数据
    title = book.get_metadata('DC', 'title')
    author = book.get_metadata('DC', 'creator')
    publisher = book.get_metadata('DC', 'publisher')
    isbn = book.get_metadata('DC', 'identifier')
    language = book.get_metadata('DC', 'language')
    
    document.title = title[0][0] if title else file_path.stem
    document.metadata['author'] = author[0][0] if author else 'Unknown'
    document.metadata['publisher'] = publisher[0][0] if publisher else ''
    document.metadata['isbn'] = isbn[0][0] if isbn else f'translated-{document.title.replace(" ", "-")}'
    document.metadata['language'] = language[0][0] if language else 'en'
    
    # 提取 CSS 样式
    for item in book.get_items():
        if item.get_type() == 8:  # CSS
            try:
                document.css += item.get_content().decode('utf-8', errors='ignore') + "\n"
            except:
                pass
    
    # 提取所有图片
    image_items = []
    for item in book.get_items():
        # 类型 1 = IMAGE, 或检查 media_type
        if item.get_type() == 1 or item.media_type.startswith('image/'):
            try:
                href = item.get_name()
                content = item.get_content()
                media_type = item.media_type
                
                # 保存图片到输出目录
                images_dir = output_dir / 'images'
                images_dir.mkdir(exist_ok=True)
                image_path = images_dir / href.split('/')[-1]
                image_path.write_bytes(content)
                
                image_info = {
                    'href': href,
                    'path': str(image_path),
                    'media_type': media_type,
                    'size': len(content)
                }
                document.images.append(image_info)
                image_items.append(image_info)
            except Exception as e:
                print(f"图片提取失败 {item.get_name()}: {e}")
    
    # 按章节顺序提取内容
    chapter_num = 0
    full_text_pos = 0
    
    # 获取 spine 顺序（阅读顺序）
    spine_items = book.spine
    
    for item in book.get_items():
        if item.get_type() == 9:  # XHTML content
            try:
                content = item.get_content().decode('utf-8', errors='ignore')
            except:
                continue
            
            # 提取纯文本（用于翻译）
            text = re.sub(r'<[^>]+>', '', content)
            
            # 提取章节标题（多级匹配）
            chapter_title = None
            
            # 1. 尝试 <title> 标签
            title_match = re.search(r'<title[^>]*>([^<]+)</title>', content)
            if title_match:
                chapter_title = title_match.group(1).strip()
            
            # 2. 尝试 <h1>, <h2> 标签
            if not chapter_title:
                heading_match = re.search(r'<h[12][^>]*>([^<]+)</h[12]>', content, re.IGNORECASE)
                if heading_match:
                    chapter_title = heading_match.group(1).strip()
            
            # 3. 使用文件名
            if not chapter_title:
                file_name = item.get_name()
                # 从文件名提取（如 Text/Chapter01.xhtml -> Chapter 1）
                name_match = re.search(r'(?:Chapter|Chap|Ch)\.?(\d+)', file_name, re.IGNORECASE)
                if name_match:
                    chapter_title = f"第{name_match.group(1)}章"
                elif 'Cover' in file_name:
                    chapter_title = '封面'
                elif 'HalfTitle' in file_name:
                    chapter_title = '半标题'
                elif 'TitlePage' in file_name:
                    chapter_title = '标题页'
                elif 'Copyright' in file_name:
                    chapter_title = '版权页'
                elif 'Dedication' in file_name:
                    chapter_title = '献词'
                elif 'Acknowledgements' in file_name:
                    chapter_title = '致谢'
                elif 'Introduction' in file_name:
                    chapter_title = '引言'
                elif 'Notes' in file_name:
                    chapter_title = '注释'
                elif 'Bibliography' in file_name:
                    chapter_title = '参考文献'
                else:
                    chapter_title = Path(file_name).stem
            
            chapter_title = chapter_title or f"第{chapter_num + 1}章"
            
            # 提取本章图片引用
            chapter_images = []
            # 匹配 <img src="..."> 和 SVG <image xlink:href="..."> / <image href="...">
            img_patterns = [
                r'<img[^>]+src=["\']([^"\']+)["\']',
                r'<image[^>]+xlink:href=["\']([^"\']+)["\']',
                r'<image[^>]+href=["\']([^"\']+)["\']',
            ]
            
            for pattern in img_patterns:
                for img_match in re.finditer(pattern, content):
                    img_href = img_match.group(1)
                    # 规范化路径（移除 ../ 前缀）
                    img_href_normalized = img_href.replace('../', '').replace('./', '')
                    
                    # 找到对应的图片信息
                    for img_info in image_items:
                        img_info_href = img_info['href'].replace('../', '').replace('./', '')
                        if (img_info_href == img_href_normalized or 
                            img_info_href.endswith(img_href_normalized) or
                            Path(img_info_href).name == Path(img_href_normalized).name):
                            if img_info not in chapter_images:  # 避免重复
                                chapter_images.append(img_info)
                            break
            
            # 保留有文本或有图片的章节
            if text.strip() or chapter_images:
                chapter = Chapter(
                    id=chapter_num,
                    title=chapter_title,
                    content=content,
                    text=text if text.strip() else f"（{chapter_title}）",
                    images=chapter_images,
                    start_pos=full_text_pos,
                    end_pos=full_text_pos + len(text)
                )
                document.chapters.append(chapter)
                full_text_pos += len(text)
                chapter_num += 1
    
    # 保存图片索引
    if document.images:
        images_index = output_dir / 'images_index.json'
        with open(images_index, 'w', encoding='utf-8') as f:
            json.dump(document.images, f, ensure_ascii=False, indent=2)
    
    document.metadata["total_chapters"] = chapter_num
    document.metadata["total_images"] = len(document.images)
    document.metadata["total_chars"] = full_text_pos
    
    # 保存文档结构（用于后续合并）
    try:
        from modules.structure import create_structure_from_document, StructureManager
        structure = create_structure_from_document(document, str(output_dir))
        structure_mgr = StructureManager(str(output_dir))
        structure_mgr.save_structure(structure)
    except Exception as e:
        print(f"⚠️ 结构保存失败：{e}")
    
    print(f"✅ EPUB 解析完成：{chapter_num}章，{len(document.images)}张图片，{full_text_pos}字符")
    
    return document


async def _parse_pdf(file_path: Path, output_dir: Path) -> Document:
    """解析 PDF - 保留页面结构"""
    doc = fitz.open(file_path)
    document = Document(title=file_path.stem, source_type='pdf')
    
    full_text_pos = 0
    chapter_num = 0
    
    for page_num, page in enumerate(doc):
        # 提取文本（带格式信息）
        text = page.get_text("text")
        
        # 提取图片
        page_images = []
        for img in page.get_images(full=True):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                if base_image:
                    image_ext = base_image["ext"]
                    image_data = base_image["image"]
                    
                    # 保存图片
                    images_dir = output_dir / 'images'
                    images_dir.mkdir(exist_ok=True)
                    image_path = images_dir / f"page{page_num+1}_img{len(page_images)+1}.{image_ext}"
                    image_path.write_bytes(image_data)
                    
                    page_images.append({
                        'page': page_num,
                        'path': str(image_path),
                        'media_type': f'image/{image_ext}',
                        'size': len(image_data)
                    })
                    document.images.append(page_images[-1])
            except Exception as e:
                print(f"PDF 图片提取失败：{e}")
        
        if text.strip():
            chapter = Chapter(
                id=chapter_num,
                title=f"第{page_num + 1}页",
                content=text,
                text=text,
                images=page_images,
                start_pos=full_text_pos,
                end_pos=full_text_pos + len(text)
            )
            document.chapters.append(chapter)
            full_text_pos += len(text)
            chapter_num += 1
    
    doc.close()
    
    document.metadata["total_pages"] = len(doc)
    document.metadata["total_images"] = len(document.images)
    document.metadata["total_chars"] = full_text_pos
    
    print(f"✅ PDF 解析完成：{chapter_num}页，{len(document.images)}张图片，{full_text_pos}字符")
    
    return document


async def _parse_fb2(file_path: Path, output_dir: Path) -> Document:
    """解析 FB2 (XML 格式)"""
    import xml.etree.ElementTree as ET
    
    tree = ET.parse(file_path)
    root = tree.getroot()
    
    # 命名空间
    ns = {'fb': 'http://www.gribuser.ru/xml/fictionbook/2.0'}
    
    # 提取标题
    title_elem = root.find('.//fb:book-title', ns)
    title = title_elem.text if title_elem is not None else file_path.stem
    
    document = Document(title=title, source_type='fb2')
    
    # 提取正文（按 section 分章）
    body = root.find('.//fb:body', ns)
    if body is not None:
        full_text_pos = 0
        chapter_num = 0
        
        for section in body.findall('.//fb:section', ns):
            section_title_elem = section.find('fb:title', ns)
            section_title = section_title_elem.text if section_title_elem is not None else f"第{chapter_num + 1}章"
            
            section_text = ""
            for p in section.findall('.//fb:p', ns):
                if p.text:
                    section_text += p.text.strip() + "\n"
            
            if section_text.strip():
                chapter = Chapter(
                    id=chapter_num,
                    title=section_title,
                    content=section_text,
                    text=section_text,
                    start_pos=full_text_pos,
                    end_pos=full_text_pos + len(section_text)
                )
                document.chapters.append(chapter)
                full_text_pos += len(section_text)
                chapter_num += 1
        
        document.metadata["total_chapters"] = chapter_num
        document.metadata["total_chars"] = full_text_pos
    
    return document


async def _parse_txt(file_path: Path, output_dir: Path) -> Document:
    """解析 TXT"""
    import chardet
    
    # 自动检测编码
    raw = file_path.read_bytes()
    detected = chardet.detect(raw)
    encoding = detected['encoding'] or 'utf-8'
    
    try:
        text = file_path.read_text(encoding=encoding)
    except:
        text = file_path.read_text(encoding='utf-8', errors='ignore')
    
    # 尝试按章节分割（基于常见章节标记）
    chapters = []
    chapter_pattern = r'(?:^|\n)\s*(?:第 [零一二三四五六七八九十百\d]+章|Chapter\s+\d+|PART\s+[IVX]+)\s*[:：]?\s*(.*?)\s*(?=\n|$)'
    
    matches = list(re.finditer(chapter_pattern, text, re.IGNORECASE))
    
    if matches:
        # 有明确章节标记
        full_text_pos = 0
        for i, match in enumerate(matches):
            start = match.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            chapter_text = text[start:end]
            
            chapter_title = match.group(0).strip()
            
            chapter = Chapter(
                id=i,
                title=chapter_title,
                content=chapter_text,
                text=chapter_text,
                start_pos=full_text_pos,
                end_pos=full_text_pos + len(chapter_text)
            )
            chapters.append(chapter)
            full_text_pos += len(chapter_text)
    else:
        # 无章节标记，作为单章
        chapters = [Chapter(
            id=0,
            title=file_path.stem,
            content=text,
            text=text,
            start_pos=0,
            end_pos=len(text)
        )]
    
    document = Document(
        title=file_path.stem,
        source_type='txt',
        chapters=chapters,
        metadata={"total_chars": len(text)}
    )
    
    return document


async def _parse_docx(file_path: Path, output_dir: Path) -> Document:
    """解析 DOCX"""
    from docx import Document as DocxDocument
    
    doc = DocxDocument(file_path)
    
    # 按段落分组（尝试识别章节）
    chapters = []
    current_chapter = []
    chapter_num = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查是否是标题（样式或内容）
        is_heading = para.style.name.startswith('Heading') if para.style else False
        is_chapter_title = re.match(r'^(?:第 [零一二三四五六七八九十百\d]+章|Chapter\s+\d+)', text)
        
        if is_heading or is_chapter_title:
            # 新章节开始
            if current_chapter:
                chapter_text = "\n\n".join(current_chapter)
                chapters.append(Chapter(
                    id=chapter_num,
                    title=current_chapter[0] if current_chapter else f"第{chapter_num + 1}章",
                    content=chapter_text,
                    text=chapter_text,
                    start_pos=0,
                    end_pos=len(chapter_text)
                ))
                chapter_num += 1
            current_chapter = [text]
        else:
            current_chapter.append(text)
    
    # 最后一章
    if current_chapter:
        chapter_text = "\n\n".join(current_chapter)
        chapters.append(Chapter(
            id=chapter_num,
            title=current_chapter[0] if current_chapter else f"第{chapter_num + 1}章",
            content=chapter_text,
            text=chapter_text,
            start_pos=0,
            end_pos=len(chapter_text)
        ))
    
    # 标题
    title = doc.core_properties.title or file_path.stem
    
    document = Document(
        title=title,
        source_type='docx',
        chapters=chapters if chapters else [Chapter(
            id=0,
            title=title,
            content="\n\n".join([p.text for p in doc.paragraphs if p.text.strip()]),
            text="\n\n".join([p.text for p in doc.paragraphs if p.text.strip()]),
            start_pos=0,
            end_pos=0
        )],
        metadata={"total_chars": sum(len(c.text) for c in chapters)}
    )
    
    return document


async def _parse_html(file_path: Path, output_dir: Path) -> Document:
    """解析 HTML"""
    from bs4 import BeautifulSoup
    
    html = file_path.read_text(encoding='utf-8', errors='ignore')
    soup = BeautifulSoup(html, 'html.parser')
    
    # 清理不需要的标签
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    
    # 提取标题
    title_elem = soup.find('title')
    title = title_elem.string if title_elem else file_path.stem
    
    # 尝试找正文
    main = soup.find('main') or soup.find('article') or soup.find('body') or soup
    
    # 提取章节（按 h1/h2 分割）
    chapters = []
    chapter_num = 0
    current_content = []
    current_title = title
    
    for element in main.children:
        if element.name in ['h1', 'h2']:
            # 新章节
            if current_content:
                chapter_text = "\n\n".join(current_content)
                chapters.append(Chapter(
                    id=chapter_num,
                    title=current_title,
                    content=chapter_text,
                    text=chapter_text,
                    start_pos=0,
                    end_pos=len(chapter_text)
                ))
                chapter_num += 1
                current_content = []
            current_title = element.get_text().strip()
        elif element.name and element.get_text(strip=True):
            current_content.append(element.get_text(separator='\n'))
    
    # 最后一章
    if current_content:
        chapter_text = "\n\n".join(current_content)
        chapters.append(Chapter(
            id=chapter_num,
            title=current_title,
            content=chapter_text,
            text=chapter_text,
            start_pos=0,
            end_pos=len(chapter_text)
        ))
    
    document = Document(
        title=title,
        source_type='html',
        chapters=chapters if chapters else [Chapter(
            id=0,
            title=title,
            content=main.get_text(separator='\n\n'),
            text=main.get_text(separator='\n\n'),
            start_pos=0,
            end_pos=len(main.get_text(separator='\n\n'))
        )],
        metadata={"total_chars": sum(len(c.text) for c in chapters)}
    )
    
    return document


# 解析器映射
PARSER_MAP = {
    'pdf': _parse_pdf,
    'epub': _parse_epub,
    'fb2': _parse_fb2,
    'txt': _parse_txt,
    'docx': _parse_docx,
    'html': _parse_html,
}
